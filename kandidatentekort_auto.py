#!/usr/bin/env python3
"""
KANDIDATENTEKORT.NL - WEBHOOK AUTOMATION V5.2
Deploy: Render.com | Updated: 2026-04-10
- V2: Pipedrive organization, person, deal creation
- V3: Claude AI vacancy analysis + report email
- V3.1: Professional report template with Before/After comparison
- V3.2: PDF, DOCX and Word file extraction for vacancy analysis
- V3.3: Fixed Typeform file download with authentication
- V4.0: ULTIMATE email template - Score visualization, Category breakdown,
        Before/After comparison, Full improved text, Numbered checklist, Bonus tips
- V4.1: OUTLOOK COMPATIBLE - Full table-based layout, MSO conditionals, no flex/gradients
- V5.1: TRUST-FIRST EMAIL NURTURE - 8 automated follow-up emails over 30 days
"""

import os
import io
import re
import json
import logging
import smtplib
import requests
import threading
import time
import hmac
import hashlib
import base64
from datetime import datetime, timedelta
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from flask import Flask, request, jsonify
from markupsafe import escape as html_escape

# Setup logging FIRST before any logger usage
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Report generator (V5.3: Figma report templates + Supabase Storage)
REPORT_BUILDER_AVAILABLE = False
try:
    from generator.report_builder import build_hosted_rapport, build_email_summary
    from generator.storage_uploader import upload_rapport, upload_analysis_json, fetch_analysis_json
    REPORT_BUILDER_AVAILABLE = True
    logger.info("✅ Report builder loaded")
except ImportError as e:
    logger.warning(f"⚠️ Report builder import failed: {e}")

# PDF and DOCX extraction
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

app = Flask(__name__)

# Config
ANTHROPIC_API_KEY = os.getenv('ANTHROPIC_API_KEY')
PIPEDRIVE_API_TOKEN = os.getenv('PIPEDRIVE_API_TOKEN')
TYPEFORM_API_TOKEN = os.getenv('TYPEFORM_API_TOKEN')  # For file downloads
RESEND_API_KEY = os.getenv('RESEND_API_KEY')
RESEND_FROM_EMAIL = os.getenv('RESEND_FROM_EMAIL', 'noreply@kandidatentekort.nl')
GMAIL_USER = os.getenv('GMAIL_USER', 'artsrecruitin@gmail.com')
GMAIL_APP_PASSWORD = os.getenv('GMAIL_APP_PASSWORD') or os.getenv('GMAIL_PASS')
PIPEDRIVE_BASE = "https://api.pipedrive.com/v1"
PIPELINE_ID = 4
STAGE_ID = 21
ADMIN_SECRET = os.getenv('ADMIN_SECRET', '')  # Required for test/debug/nurture endpoints

# Email Nurture Custom Field Keys (from Pipedrive)
FIELD_RAPPORT_VERZONDEN = "337f9ccca15334e6e4f937ca5ef0055f13ed0c63"
FIELD_EMAIL_SEQUENCE_STATUS = "22d33c7f119119e178f391a272739c571cf2e29b"
FIELD_LAATSTE_EMAIL = "753f37a1abc8e161c7982c1379a306b21fae1bab"
FIELD_ANALYSE_STORAGE_PREFIX = os.getenv("FIELD_ANALYSE_STORAGE_PREFIX", "")  # Pipedrive text field for Supabase path

# Email sequence timing (days after rapport verzonden)
# Emails 1-3: VALUE DRIP (verbeterde tekst → marktanalyse → kanaalstrategie)
# Emails 4-8: TRUST-BUILDING (tips + check-ins)
EMAIL_SCHEDULE = {
    1: {"day": 1, "name": "Drip: Verbeterde Vacaturetekst"},
    2: {"day": 3, "name": "Drip: Marktanalyse + Salaris"},
    3: {"day": 5, "name": "Drip: Kanaalstrategie + Actieplan"},
    4: {"day": 8, "name": "Tip Functietitel"},
    5: {"day": 11, "name": "Tip Salaris"},
    6: {"day": 14, "name": "Tip Opening"},
    7: {"day": 21, "name": "Gesprek Aanbod"},
    8: {"day": 30, "name": "Final Check-in"},
}

# Stage filter: Only send nurture emails to deals in stage 21 (Gekwalificeerd)
# Deals in stage 22+ have active contact, so no automated emails needed
NURTURE_ACTIVE_STAGE = 21  # Gekwalificeerd


def extract_text_from_file(file_url):
    """
    Download and extract text from PDF, DOCX, or DOC files.
    Returns extracted text or empty string on failure.
    Typeform file URLs require Bearer token authentication.
    """
    if not file_url:
        return ""

    try:
        logger.info(f"📄 Downloading file: {file_url[:80]}...")

        # Prepare headers - Typeform API requires authentication
        headers = {}
        if TYPEFORM_API_TOKEN and 'typeform.com' in file_url:
            headers['Authorization'] = f'Bearer {TYPEFORM_API_TOKEN}'
            logger.info("🔑 Using Typeform API authentication")

        # Download the file
        response = requests.get(file_url, headers=headers, timeout=30)

        # Log response details for debugging
        content_type = response.headers.get('content-type', 'unknown')
        logger.info(f"📦 Response: status={response.status_code}, content-type={content_type}, size={len(response.content)} bytes")

        if response.status_code != 200:
            logger.error(f"❌ Failed to download file: {response.status_code}")
            return ""

        content = response.content

        # Check if we got an error page instead of the file
        if len(content) < 100 and b'error' in content.lower():
            logger.error(f"❌ Got error response: {content[:200]}")
            return ""

        # Detect file type by content (magic bytes) - most reliable
        if content[:4] == b'%PDF':
            logger.info("📄 Detected PDF by magic bytes")
            return extract_pdf_text(content)
        elif content[:2] == b'PK':  # DOCX/XLSX/ZIP files start with PK
            logger.info("📄 Detected DOCX/ZIP by magic bytes")
            return extract_docx_text(content)
        elif content[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':  # Old .doc format (OLE)
            logger.warning("⚠️ Old .doc (OLE) format - not supported, try converting to DOCX")
            return ""

        # Fallback: try by content-type header
        if 'pdf' in content_type:
            return extract_pdf_text(content)
        elif 'wordprocessingml' in content_type or 'msword' in content_type:
            return extract_docx_text(content)

        # Last resort: try by URL extension
        file_url_lower = file_url.lower()
        if '.pdf' in file_url_lower:
            return extract_pdf_text(content)
        elif '.docx' in file_url_lower:
            return extract_docx_text(content)

        logger.warning(f"⚠️ Could not determine file type. Content starts with: {content[:20]}")
        return ""

    except Exception as e:
        logger.error(f"❌ File extraction error: {e}")
        return ""


def extract_pdf_text(content):
    """Extract text from PDF content"""
    if not PDF_AVAILABLE:
        logger.error("❌ PyPDF2 not available")
        return ""

    try:
        pdf_file = io.BytesIO(content)
        reader = PdfReader(pdf_file)

        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        full_text = "\n".join(text_parts)
        logger.info(f"✅ PDF extracted: {len(full_text)} characters from {len(reader.pages)} pages")
        return full_text.strip()

    except Exception as e:
        logger.error(f"❌ PDF extraction failed: {e}")
        return ""


def extract_docx_text(content):
    """Extract text from DOCX content"""
    if not DOCX_AVAILABLE:
        logger.error("❌ python-docx not available")
        return ""

    try:
        docx_file = io.BytesIO(content)
        doc = Document(docx_file)

        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        full_text = "\n".join(text_parts)
        logger.info(f"✅ DOCX extracted: {len(full_text)} characters")
        return full_text.strip()

    except Exception as e:
        logger.error(f"❌ DOCX extraction failed: {e}")
        return ""


def get_confirmation_email_html(voornaam, bedrijf, functie):
    return f'''<!DOCTYPE html><html><head><meta charset="UTF-8"></head>
<body style="margin:0;padding:0;font-family:Inter,-apple-system,sans-serif;background:#f9fafb;">
<table width="100%" style="padding:40px 20px;"><tr><td align="center">
<table width="600" style="background:#fff;border-radius:12px;box-shadow:0 8px 32px rgba(44,62,80,0.12);">
<tr><td style="background:linear-gradient(135deg,#ff6b35,#e55a2b);color:#fff;padding:40px 30px;text-align:center;">
<div style="font-size:28px;font-weight:800;">✅ Ontvangen!</div>
<div style="font-size:16px;opacity:0.95;">Je vacature-analyse aanvraag is binnen</div></td></tr>
<tr><td style="padding:35px 30px;">
<p style="font-size:19px;font-weight:700;">Hoi {voornaam},</p>
<p style="color:#374151;">Bedankt! We hebben je vacature voor <strong style="color:#ff6b35;">{functie}</strong> bij <strong style="color:#ff6b35;">{bedrijf}</strong> ontvangen.</p>
<table width="100%" style="background:#f0f4f8;border-left:5px solid #ff6b35;border-radius:0 12px 12px 0;margin:25px 0;">
<tr><td style="padding:25px;">
<div style="font-size:18px;font-weight:700;color:#2c3e50;">⏰ Wat kun je verwachten?</div>
<ul style="color:#374151;padding-left:20px;">
<li><strong>Binnen 24 uur</strong> ontvang je een uitgebreide analyse</li>
<li>Concrete verbeterpunten voor meer sollicitanten</li>
<li>Score-overzicht op 6 belangrijke gebieden</li>
<li>Direct toepasbare tips</li></ul></td></tr></table>
<p style="color:#374151;">Vragen? Reply gewoon op deze email.</p></td></tr>
<tr><td style="padding:0 30px 35px;border-top:1px solid #f1f3f4;">
<table style="padding-top:25px;"><tr><td>
<p style="margin:0 0 5px;font-weight:700;color:#2c3e50;">Wouter Arts</p>
<p style="margin:0;color:#6b7280;font-size:14px;">Founder & Recruitment Specialist</p>
<p style="margin:0;color:#ff6b35;font-size:14px;font-weight:600;">Kandidatentekort.nl</p>
</td></tr></table></td></tr>
<tr><td style="background:#2c3e50;color:#fff;padding:20px 30px;text-align:center;font-size:12px;">
© 2025 Kandidatentekort.nl | Recruitin B.V.</td></tr>
</table></td></tr></table></body></html>'''


def send_email(to_email, subject, html_body):
    """Send email via Resend API (production-ready transactional email)"""
    logger.info(f"📧 Sending via Resend to: {to_email}")

    if not RESEND_API_KEY:
        logger.error("❌ RESEND_API_KEY not set!")
        return False

    try:
        response = requests.post(
            "https://api.resend.com/emails",
            headers={"Authorization": f"Bearer {RESEND_API_KEY}"},
            json={
                "from": RESEND_FROM_EMAIL,
                "to": to_email,
                "subject": subject,
                "html": html_body
            },
            timeout=10
        )

        if response.status_code == 200:
            logger.info(f"✅ Email sent via Resend to {to_email}")
            return True
        else:
            logger.error(f"❌ Resend API error: {response.status_code} - {response.text}")
            return False
    except Exception as e:
        logger.error(f"❌ Email failed: {e}")
        return False


def send_confirmation_email(to_email, voornaam, bedrijf, functie):
    return send_email(to_email, f"✅ Ontvangen: Vacature-analyse voor {functie}",
                      get_confirmation_email_html(voornaam, bedrijf, functie))


def analyze_vacancy_with_claude(vacature_text, bedrijf, sector=""):
    """Analyze vacancy text with Claude AI and return structured analysis"""
    if not ANTHROPIC_API_KEY:
        logger.error("❌ ANTHROPIC_API_KEY not set!")
        return None

    prompt = f"""Je bent een senior recruitment copywriter met 15+ jaar ervaring in technische en industriële vacatures in Nederland. Analyseer deze vacaturetekst grondig en herschrijf hem zodat hij MEER sollicitaties oplevert.

Return ONLY valid JSON, geen extra tekst.

=== VACATURETEKST ===
{vacature_text[:3000]}

=== CONTEXT ===
Bedrijf: {bedrijf}
Sector: {sector or 'onbekend'}

=== JSON STRUCTUUR (volg EXACT) ===
{{
    "overall_score": 64,
    "samenvatting": "Directe samenvatting in 2-3 zinnen. Noem de score, het sterkste punt, en het kritiekste verbeterpunt met concrete impact (bijv. 'Door het ontbreken van salarisindicatie mis je ~35% potentiële sollicitanten').",
    "score_section": "Titel: 7/10 - Omschrijving: 7/10 - Salaris: 4/10 - Branding: 5/10",
    "categories": [
        {{"name": "Vacaturetitel & Vindbaarheid", "score": 75, "status": "ok"}},
        {{"name": "Functieomschrijving", "score": 70, "status": "ok"}},
        {{"name": "Salaris & Arbeidsvoorwaarden", "score": 35, "status": "bad"}},
        {{"name": "Employer Branding", "score": 40, "status": "bad"}},
        {{"name": "Kandidaat Experience", "score": 65, "status": "warning"}},
        {{"name": "Kanaalstrategie", "score": 55, "status": "warning"}},
        {{"name": "Concurrentiekracht", "score": 50, "status": "warning"}},
        {{"name": "SEO & Online Vindbaarheid", "score": 80, "status": "ok"}}
    ],
    "market_analysis": {{
        "competing_vacancies": 23,
        "potential_candidates": 142,
        "market_median_salary": "€4.800",
        "supply_demand_ratio": "3.2x"
    }},
    "salary_benchmark": {{
        "offered_range": "€4.200 - €5.500",
        "market_range": "€4.800 - €6.200",
        "difference": "-12% onder markt",
        "warning": "Salaris ligt onder marktgemiddelde"
    }},
    "top_3_improvements": ["Concrete verbetering met verwachte impact", "...", "..."],
    "improved_text": "Volledige herschreven vacaturetekst — zie instructies hieronder",
    "action_items": ["Concrete, uitvoerbare actie met deadline-suggestie", "...", "...", "...", "..."],
    "recommended_channels": [
        {{"name": "Indeed", "description": "Bereik X kandidaten in deze regio/sector", "status": "AANBEVOLEN"}},
        {{"name": "LinkedIn Jobs", "description": "Gericht op ervaren professionals in {sector or 'deze sector'}", "status": "AANBEVOLEN"}}
    ],
    "bonus_tips": ["Sectorspecifieke tip", "Tip op basis van huidige arbeidsmarkt"]
}}

=== INSTRUCTIES VOOR improved_text ===
Dit is het BELANGRIJKSTE veld. De herschreven tekst moet direct bruikbaar zijn.

VERPLICHTE STRUCTUUR van improved_text:
1. PAKKENDE OPENING (1-2 zinnen): Begin met een vraag of bold statement gericht aan de kandidaat. NOOIT beginnen met "Wij zoeken" of "Voor onze opdrachtgever".
2. OVER DE ROL (3-5 bullet points): Wat ga je DOEN, niet wat je MOET KUNNEN. Gebruik actieve werkwoorden.
3. WAT JE MEEBRENGT (3-5 bullet points): Harde eisen vs. nice-to-haves gescheiden.
4. WAT JE KRIJGT (4-6 bullet points): Concreet — euro's, dagen, mogelijkheden. NOOIT "marktconform" of "passend salaris".
5. OVER HET BEDRIJF (2-3 zinnen): Specifiek over {bedrijf} — omvang, cultuur, projecten. Gebruik info uit de originele tekst.
6. SOLLICITEER-CTA (1-2 zinnen): Laagdrempelig, met naam contactpersoon als die in de tekst staat.

STIJLREGELS voor improved_text:
- BEHOUD de toon van het origineel (formeel bedrijf = formeel, informeel = informeel)
- Gebruik SPECIFIEKE details uit de originele tekst (projecten, producten, klanten, tools)
- Als de originele tekst geen salaris noemt: voeg een realistische indicatie toe op basis van functie/sector/regio
- VERBODEN woorden/frases: "dynamisch team", "no-nonsense", "familiaire sfeer", "korte lijnen", "marktconform", "uitdagende functie", "enthousiaste collega's", "passend salaris"
- Schrijf alsof je de kandidaat persoonlijk aanspreekt
- Minimaal 250 woorden, maximaal 500 woorden

=== SCORING REGELS ===
- Scores 0-100. overall_score = gemiddelde van 8 categorieën
- status: "ok" (>=65), "warning" (45-64), "bad" (<45)
- Wees EERLIJK en KRITISCH. De meeste vacatures scoren 40-65. Score >75 alleen als de tekst écht bovengemiddeld is.
- market_analysis: schat REALISTISCH in op basis van sector, functie, en Nederlandse arbeidsmarkt (niet te optimistisch)
- salary_benchmark: als geen salaris vermeld, schat in op basis van functie/sector. Vermeld dit als warning.
- action_items: 5 concrete stappen die de klant VANDAAG kan uitvoeren, in volgorde van impact
- Alles in het Nederlands"""

    try:
        logger.info("🤖 Starting Claude analysis...")
        r = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "x-api-key": ANTHROPIC_API_KEY,
                "anthropic-version": "2023-06-01",
                "content-type": "application/json"
            },
            json={
                "model": "claude-sonnet-4-6",
                "max_tokens": 6000,
                "messages": [{"role": "user", "content": prompt}]
            },
            timeout=120
        )

        if r.status_code == 200:
            response_text = r.json()['content'][0]['text']
            # Extract JSON from response
            json_start = response_text.find('{')
            json_end = response_text.rfind('}') + 1
            if json_start >= 0 and json_end > json_start:
                analysis = json.loads(response_text[json_start:json_end])
                logger.info(f"✅ Claude analysis complete: score={analysis.get('overall_score')}")
                return analysis
            else:
                logger.error(f"❌ No JSON found in Claude response")
                return None
        else:
            logger.error(f"❌ Claude API error: {r.status_code} - {r.text[:200]}")
            return None

    except Exception as e:
        logger.error(f"❌ Claude analysis failed: {e}")
        return None


def get_analysis_email_html(voornaam, bedrijf, analysis, original_text=""):
    """
    Generate the ULTIMATE professional analysis report email HTML V4.1
    OUTLOOK COMPATIBLE - No flex, no gradients, all table-based layout
    Features: Score visualization, Before/After, Checklist, Improved text, Tips
    """
    score = analysis.get('overall_score', 'N/A')
    score_section = analysis.get('score_section', '')
    improvements = analysis.get('top_3_improvements', [])
    improved_text = analysis.get('improved_text', '')
    bonus_tips = analysis.get('bonus_tips', [])

    # Generate HTML for improvements (numbered) - OUTLOOK COMPATIBLE
    improvements_html = ''.join([
        f'''<tr>
        <td style="padding:12px 0;border-bottom:1px solid #FDE68A;">
        <table cellpadding="0" cellspacing="0" width="100%"><tr>
        <td width="40" valign="top"><table cellpadding="0" cellspacing="0"><tr><td style="width:32px;height:32px;background-color:#F59E0B;text-align:center;font-weight:bold;color:white;font-size:14px;font-family:Arial,sans-serif;mso-line-height-rule:exactly;line-height:32px;">{i+1}</td></tr></table></td>
        <td style="padding-left:12px;color:#78350F;font-size:14px;line-height:22px;font-family:Arial,sans-serif;">{imp}</td>
        </tr></table>
        </td></tr>''' for i, imp in enumerate(improvements)
    ])

    # Generate HTML for bonus tips - OUTLOOK COMPATIBLE (table-based)
    tips_html = ''.join([
        f'''<tr><td style="padding:8px 0;">
        <table cellpadding="0" cellspacing="0" width="100%" style="background-color:#ffffff;"><tr>
        <td width="40" valign="top" style="padding:12px;font-size:18px;">💡</td>
        <td style="padding:12px;color:#5B21B6;font-size:14px;line-height:22px;font-family:Arial,sans-serif;">{tip}</td>
        </tr></table>
        </td></tr>''' for tip in bonus_tips
    ])

    # Truncate original text for before/after display
    original_display = original_text[:500] + '...' if len(original_text) > 500 else original_text
    improved_preview = improved_text[:500] + '...' if len(improved_text) > 500 else improved_text

    # Escape newlines for HTML display
    original_display = original_display.replace('\n', '<br>')
    improved_preview = improved_preview.replace('\n', '<br>')
    improved_text_html = improved_text.replace('\n', '<br>')

    # Calculate score color and label based on score (0-100 scale)
    if isinstance(score, (int, float)):
        score_num = float(score)
        if score_num >= 75:
            score_color = "#10B981"
            score_bg = "#ECFDF5"
            score_border = "#10B981"
            score_label = "Uitstekend"
            score_emoji = "🏆"
        elif score_num >= 60:
            score_color = "#3B82F6"
            score_bg = "#EFF6FF"
            score_border = "#3B82F6"
            score_label = "Goed"
            score_emoji = "👍"
        elif score_num >= 45:
            score_color = "#F59E0B"
            score_bg = "#FFFBEB"
            score_border = "#F59E0B"
            score_label = "Kan beter"
            score_emoji = "📈"
        else:
            score_color = "#EF4444"
            score_bg = "#FEF2F2"
            score_border = "#EF4444"
            score_label = "Verbetering nodig"
            score_emoji = "⚠️"
    else:
        score_color = "#6B7280"
        score_bg = "#F9FAFB"
        score_border = "#6B7280"
        score_label = "Beoordeeld"
        score_emoji = "📊"

    # Parse score_section into categories - OUTLOOK COMPATIBLE
    categories_html = ""
    if score_section:
        score_parts = re.findall(r'([A-Za-z-]+):\s*(\d+)/10', score_section)
        if score_parts:
            categories_html = '<table width="100%" cellpadding="0" cellspacing="0" style="margin-top:20px;"><tr>'
            for name, cat_score in score_parts[:4]:
                cat_score_int = int(cat_score)
                if cat_score_int >= 7:
                    cat_color = "#10B981"
                    cat_icon = "✅"
                elif cat_score_int >= 5:
                    cat_color = "#F59E0B"
                    cat_icon = "⚡"
                else:
                    cat_color = "#EF4444"
                    cat_icon = "❗"
                categories_html += f'''<td width="25%" align="center" style="padding:10px;">
                <table cellpadding="0" cellspacing="0"><tr><td align="center" style="font-size:24px;padding-bottom:4px;">{cat_icon}</td></tr>
                <tr><td align="center" style="font-size:24px;font-weight:bold;color:{cat_color};font-family:Arial,sans-serif;">{cat_score}</td></tr>
                <tr><td align="center" style="font-size:11px;color:#6B7280;text-transform:uppercase;font-family:Arial,sans-serif;">{name}</td></tr></table>
                </td>'''
            categories_html += '</tr></table>'

    return f'''<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1.0">
<meta http-equiv="X-UA-Compatible" content="IE=edge">
<!--[if mso]>
<xml>
<o:OfficeDocumentSettings>
<o:AllowPNG/>
<o:PixelsPerInch>96</o:PixelsPerInch>
</o:OfficeDocumentSettings>
</xml>
<![endif]-->
<style type="text/css">
body, table, td {{margin:0;padding:0;font-family:Arial,Helvetica,sans-serif;}}
img {{border:0;height:auto;line-height:100%;outline:none;text-decoration:none;}}
table {{border-collapse:collapse !important;}}
</style>
</head>
<body style="margin:0;padding:0;background-color:#f3f4f6;width:100%;">

<!-- OUTLOOK WRAPPER -->
<!--[if mso]>
<table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0" style="background-color:#f3f4f6;">
<tr><td align="center" style="padding:30px 0;">
<![endif]-->

<table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0" style="background-color:#f3f4f6;">
<tr><td align="center" style="padding:30px 15px;">

<table role="presentation" width="650" cellpadding="0" cellspacing="0" border="0" style="background-color:#ffffff;max-width:650px;">

<!-- ════════════════════════════════════════════════════════════════════════ -->
<!-- HEADER -->
<!-- ════════════════════════════════════════════════════════════════════════ -->
<tr>
<td style="background-color:#1E3A8A;padding:40px 35px;">
<table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0">
<tr>
<td width="60" valign="middle">
<table role="presentation" cellpadding="0" cellspacing="0" border="0">
<tr><td style="width:56px;height:56px;background-color:#FF6B35;text-align:center;font-size:26px;font-weight:bold;color:#ffffff;font-family:Arial,sans-serif;line-height:56px;">R</td></tr>
</table>
</td>
<td style="padding-left:16px;" valign="middle">
<p style="margin:0;color:#ffffff;font-size:20px;font-weight:bold;font-family:Arial,sans-serif;">RECRUITIN</p>
<p style="margin:4px 0 0 0;color:#E0E7FF;font-size:12px;font-family:Arial,sans-serif;">Vacature Intelligence Platform</p>
</td>
</tr>
</table>
<table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0" style="margin-top:30px;">
<tr><td>
<p style="margin:0 0 8px 0;color:#93C5FD;font-size:12px;font-weight:bold;text-transform:uppercase;letter-spacing:1px;font-family:Arial,sans-serif;">AI-POWERED ANALYSE</p>
<p style="margin:0;color:#ffffff;font-size:28px;font-weight:bold;font-family:Arial,sans-serif;">📊 Vacature Analyse Rapport</p>
<p style="margin:10px 0 0 0;color:#E0E7FF;font-size:15px;font-family:Arial,sans-serif;">Gepersonaliseerd voor <strong style="color:#ffffff;">{bedrijf}</strong></p>
</td></tr>
</table>
</td>
</tr>

<!-- ════════════════════════════════════════════════════════════════════════ -->
<!-- SCORE SECTION -->
<!-- ════════════════════════════════════════════════════════════════════════ -->
<tr>
<td style="padding:45px 35px;background-color:#f8fafc;" align="center">
<!-- Score Box -->
<table role="presentation" cellpadding="0" cellspacing="0" border="0" style="margin-bottom:20px;">
<tr><td align="center" style="width:140px;height:140px;border:8px solid {score_color};background-color:#ffffff;">
<p style="margin:0;font-size:52px;font-weight:bold;color:{score_color};font-family:Arial,sans-serif;line-height:1;">{score}</p>
<p style="margin:5px 0 0 0;font-size:16px;color:#9CA3AF;font-family:Arial,sans-serif;">/10</p>
</td></tr>
</table>
<!-- Score Label -->
<table role="presentation" cellpadding="0" cellspacing="0" border="0" style="margin-bottom:15px;">
<tr><td style="background-color:{score_bg};border:2px solid {score_border};padding:10px 24px;">
<p style="margin:0;font-size:14px;font-weight:bold;color:{score_color};font-family:Arial,sans-serif;">{score_emoji} {score_label}</p>
</td></tr>
</table>
<!-- Score Breakdown -->
<p style="margin:0;color:#6B7280;font-size:13px;font-family:Arial,sans-serif;line-height:1.6;max-width:450px;">{score_section}</p>
{categories_html}
</td>
</tr>

<!-- ════════════════════════════════════════════════════════════════════════ -->
<!-- INTRO -->
<!-- ════════════════════════════════════════════════════════════════════════ -->
<tr>
<td style="padding:0 35px 30px;">
<table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0">
<tr>
<td width="4" style="background-color:#FF6B35;"></td>
<td style="padding-left:20px;">
<p style="margin:0 0 12px 0;font-size:20px;font-weight:bold;color:#1F2937;font-family:Arial,sans-serif;">Hoi {voornaam}! 👋</p>
<p style="margin:0;color:#4B5563;font-size:15px;line-height:24px;font-family:Arial,sans-serif;">Bedankt voor het uploaden van je vacature via <strong style="color:#FF6B35;">kandidatentekort.nl</strong>. Onze AI heeft je tekst grondig geanalyseerd. Hieronder vind je de complete resultaten met concrete verbeteringen.</p>
</td>
</tr>
</table>
</td>
</tr>

<!-- ════════════════════════════════════════════════════════════════════════ -->
<!-- TOP 3 VERBETERPUNTEN -->
<!-- ════════════════════════════════════════════════════════════════════════ -->
<tr>
<td style="padding:0 35px 30px;">
<table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0" style="background-color:#FFFBEB;border:2px solid #F59E0B;">
<tr><td style="padding:25px;">
<!-- Header -->
<table role="presentation" cellpadding="0" cellspacing="0" border="0" style="margin-bottom:20px;">
<tr>
<td width="48" valign="middle"><table role="presentation" cellpadding="0" cellspacing="0" border="0"><tr><td style="width:44px;height:44px;background-color:#F59E0B;text-align:center;font-size:20px;line-height:44px;">🎯</td></tr></table></td>
<td style="padding-left:14px;" valign="middle">
<p style="margin:0 0 2px 0;font-size:11px;font-weight:bold;color:#B45309;text-transform:uppercase;font-family:Arial,sans-serif;">Prioriteit</p>
<p style="margin:0;font-size:18px;font-weight:bold;color:#92400E;font-family:Arial,sans-serif;">Top 3 Verbeterpunten</p>
</td>
</tr>
</table>
<!-- List -->
<table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0">{improvements_html}</table>
</td></tr>
</table>
</td>
</tr>

<!-- ════════════════════════════════════════════════════════════════════════ -->
<!-- BEFORE / AFTER -->
<!-- ════════════════════════════════════════════════════════════════════════ -->
<tr>
<td style="padding:0 35px 30px;">
<p style="margin:0 0 20px 0;text-align:center;font-size:20px;font-weight:bold;color:#1F2937;font-family:Arial,sans-serif;">📝 Voor & Na Optimalisatie</p>
<table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0">
<tr>
<!-- BEFORE -->
<td width="48%" valign="top" style="background-color:#FEF2F2;border:2px solid #FECACA;padding:18px;">
<table role="presentation" cellpadding="0" cellspacing="0" border="0" style="margin-bottom:12px;">
<tr><td style="background-color:#EF4444;padding:5px 12px;"><p style="margin:0;font-size:11px;font-weight:bold;color:#ffffff;text-transform:uppercase;font-family:Arial,sans-serif;">❌ Origineel</p></td></tr>
</table>
<table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0" style="background-color:#ffffff;border:1px solid #FECACA;">
<tr><td style="padding:14px;font-size:12px;color:#6B7280;line-height:20px;font-family:Arial,sans-serif;">{original_display}</td></tr>
</table>
</td>
<td width="4%"></td>
<!-- AFTER -->
<td width="48%" valign="top" style="background-color:#ECFDF5;border:2px solid #A7F3D0;padding:18px;">
<table role="presentation" cellpadding="0" cellspacing="0" border="0" style="margin-bottom:12px;">
<tr><td style="background-color:#10B981;padding:5px 12px;"><p style="margin:0;font-size:11px;font-weight:bold;color:#ffffff;text-transform:uppercase;font-family:Arial,sans-serif;">✅ Geoptimaliseerd</p></td></tr>
</table>
<table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0" style="background-color:#ffffff;border:1px solid #A7F3D0;">
<tr><td style="padding:14px;font-size:12px;color:#374151;line-height:20px;font-family:Arial,sans-serif;">{improved_preview}</td></tr>
</table>
</td>
</tr>
</table>
</td>
</tr>

<!-- ════════════════════════════════════════════════════════════════════════ -->
<!-- FULL IMPROVED TEXT -->
<!-- ════════════════════════════════════════════════════════════════════════ -->
<tr>
<td style="padding:0 35px 30px;">
<table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0" style="background-color:#ECFDF5;border:2px solid #10B981;">
<tr><td style="padding:25px;">
<!-- Header -->
<table role="presentation" cellpadding="0" cellspacing="0" border="0" style="margin-bottom:20px;">
<tr>
<td width="48" valign="middle"><table role="presentation" cellpadding="0" cellspacing="0" border="0"><tr><td style="width:44px;height:44px;background-color:#10B981;text-align:center;font-size:20px;line-height:44px;">✍️</td></tr></table></td>
<td style="padding-left:14px;" valign="middle">
<p style="margin:0 0 2px 0;font-size:11px;font-weight:bold;color:#047857;text-transform:uppercase;font-family:Arial,sans-serif;">Direct te gebruiken</p>
<p style="margin:0;font-size:18px;font-weight:bold;color:#065F46;font-family:Arial,sans-serif;">Verbeterde Vacaturetekst</p>
</td>
</tr>
</table>
<!-- Text -->
<table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0" style="background-color:#ffffff;border:1px solid #A7F3D0;">
<tr><td style="padding:20px;font-size:14px;color:#374151;line-height:24px;font-family:Arial,sans-serif;">{improved_text_html}</td></tr>
</table>
<p style="margin:18px 0 0 0;text-align:center;font-size:13px;color:#059669;font-family:Arial,sans-serif;">💾 Kopieer deze tekst en plaats direct in je vacature</p>
</td></tr>
</table>
</td>
</tr>

<!-- ════════════════════════════════════════════════════════════════════════ -->
<!-- BONUS TIPS -->
<!-- ════════════════════════════════════════════════════════════════════════ -->
<tr>
<td style="padding:0 35px 30px;">
<table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0" style="background-color:#F5F3FF;border:2px solid #8B5CF6;">
<tr><td style="padding:25px;">
<!-- Header -->
<table role="presentation" cellpadding="0" cellspacing="0" border="0" style="margin-bottom:20px;">
<tr>
<td width="48" valign="middle"><table role="presentation" cellpadding="0" cellspacing="0" border="0"><tr><td style="width:44px;height:44px;background-color:#8B5CF6;text-align:center;font-size:20px;line-height:44px;">🚀</td></tr></table></td>
<td style="padding-left:14px;" valign="middle">
<p style="margin:0 0 2px 0;font-size:11px;font-weight:bold;color:#6D28D9;text-transform:uppercase;font-family:Arial,sans-serif;">Extra waarde</p>
<p style="margin:0;font-size:18px;font-weight:bold;color:#5B21B6;font-family:Arial,sans-serif;">Bonus Tips van de Expert</p>
</td>
</tr>
</table>
<!-- Tips -->
<table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0">{tips_html}</table>
</td></tr>
</table>
</td>
</tr>

<!-- ════════════════════════════════════════════════════════════════════════ -->
<!-- CTA SECTION -->
<!-- ════════════════════════════════════════════════════════════════════════ -->
<tr>
<td style="padding:0 35px 35px;">
<table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0" style="background-color:#1E3A8A;">
<tr><td style="padding:35px;text-align:center;">
<p style="margin:0 0 10px 0;font-size:22px;font-weight:bold;color:#ffffff;font-family:Arial,sans-serif;">Wil je nog meer resultaat?</p>
<p style="margin:0 0 25px 0;font-size:15px;color:#E0E7FF;line-height:22px;font-family:Arial,sans-serif;">Plan een gratis adviesgesprek van 30 minuten.<br>Bespreek je vacature met een recruitment specialist.</p>
<table role="presentation" cellpadding="0" cellspacing="0" border="0" align="center">
<tr>
<td style="padding:0 6px;">
<table role="presentation" cellpadding="0" cellspacing="0" border="0">
<tr><td style="background-color:#10B981;padding:14px 24px;">
<a href="https://calendly.com/wouter-arts-/vacature-analyse-advies" style="color:#ffffff;font-size:14px;font-weight:bold;text-decoration:none;font-family:Arial,sans-serif;">📅 Plan Adviesgesprek</a>
</td></tr>
</table>
</td>
<td style="padding:0 6px;">
<table role="presentation" cellpadding="0" cellspacing="0" border="0">
<tr><td style="background-color:#25D366;padding:14px 24px;">
<a href="https://wa.me/31614314593?text=Hoi%20Wouter,%20ik%20heb%20mijn%20vacature-analyse%20ontvangen!" style="color:#ffffff;font-size:14px;font-weight:bold;text-decoration:none;font-family:Arial,sans-serif;">💬 WhatsApp</a>
</td></tr>
</table>
</td>
</tr>
</table>
</td></tr>
</table>
</td>
</tr>

<!-- ════════════════════════════════════════════════════════════════════════ -->
<!-- FOOTER -->
<!-- ════════════════════════════════════════════════════════════════════════ -->
<tr>
<td style="background-color:#111827;padding:30px 35px;">
<table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0">
<tr>
<td valign="top">
<p style="margin:0 0 4px 0;font-size:17px;font-weight:bold;color:#ffffff;font-family:Arial,sans-serif;">Wouter Arts</p>
<p style="margin:0 0 2px 0;font-size:13px;color:#9CA3AF;font-family:Arial,sans-serif;">Founder & Recruitment Specialist</p>
<p style="margin:0;font-size:14px;font-weight:bold;color:#FF6B35;font-family:Arial,sans-serif;">Kandidatentekort.nl</p>
</td>
<td valign="top" align="right">
<p style="margin:0;font-size:12px;color:#9CA3AF;line-height:22px;font-family:Arial,sans-serif;">
📞 06-14314593<br>
📧 warts@recruitin.nl<br>
🌐 kandidatentekort.nl
</p>
</td>
</tr>
</table>
<table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0" style="margin-top:20px;border-top:1px solid #374151;">
<tr><td style="padding-top:20px;text-align:center;">
<p style="margin:0;color:#6B7280;font-size:11px;font-family:Arial,sans-serif;">© 2025 Kandidatentekort.nl | Recruitin B.V. | Made with ❤️ in Nederland</p>
</td></tr>
</table>
</td>
</tr>

</table>
</td></tr>
</table>

<!--[if mso]>
</td></tr>
</table>
<![endif]-->

</body>
</html>'''


def send_analysis_email(to_email, voornaam, bedrijf, functie, analysis, original_text=""):
    """Send the vacancy analysis report email with hosted rapport"""
    rapport_url = ""
    email_html = ""

    # V5.2: Generate hosted rapport + send email with CTA to rapport
    if REPORT_BUILDER_AVAILABLE:
        try:
            logger.info(f"📄 Building hosted rapport for {bedrijf}...")
            rapport_html = build_hosted_rapport(analysis, bedrijf, functie)

            logger.info(f"☁️ Uploading rapport to Supabase Storage...")
            rapport_url = upload_rapport(rapport_html, bedrijf)

            logger.info(f"📧 Building email summary...")
            email_html = build_email_summary(analysis, bedrijf, functie, rapport_url)

            logger.info(f"✅ Rapport generated: {rapport_url}")
        except Exception as e:
            logger.error(f"⚠️ Report builder error: {e}, falling back to legacy email")
            # Fallback to legacy email if report builder fails
            email_html = get_analysis_email_html(voornaam, bedrijf, analysis, original_text)
    else:
        # No report builder available - use legacy email
        logger.info("⚠️ Report builder not available, using legacy email")
        email_html = get_analysis_email_html(voornaam, bedrijf, analysis, original_text)

    return send_email(
        to_email,
        f"🎯 Jouw Vacature-Analyse voor {bedrijf} is Klaar!",
        email_html
    )


def parse_typeform_data(webhook_data):
    """
    Parse Typeform webhook - handles ALL field types robustly
    """
    result = {
        'email': '',
        'voornaam': 'daar',
        'contact': 'Onbekend',
        'telefoon': '',
        'bedrijf': 'Onbekend',
        'vacature': '',
        'functie': 'vacature',
        'sector': '',
        'file_url': ''
    }

    try:
        # ============================================
        # V5.1: Check for Zapier FLAT format first
        # ============================================
        if 'email' in webhook_data and 'form_response' not in webhook_data:
            logger.info("📋 Detected Zapier/Jotform FLAT format")
            result['email'] = webhook_data.get('email', '')
            result['voornaam'] = webhook_data.get('voornaam', 'daar')
            result['contact'] = f"{webhook_data.get('voornaam', '')} {webhook_data.get('achternaam', '')}".strip() or 'Onbekend'
            result['telefoon'] = webhook_data.get('telefoon', webhook_data.get('phone', ''))
            result['bedrijf'] = webhook_data.get('bedrijf', webhook_data.get('company', 'Onbekend'))
            result['vacature'] = webhook_data.get('vacature', webhook_data.get('vacancy', ''))
            result['functie'] = (webhook_data.get('functie', '') or webhook_data.get('vacature', 'vacature'))[:50]
            result['sector'] = webhook_data.get('sector', '')
            result['file_url'] = webhook_data.get('file_url', webhook_data.get('file', ''))

            # Jotform sends q{n} or q{n}_{label} fields — scan all keys for vacancy text
            if not result['vacature']:
                vacancy_keywords = ['vacature', 'tekst', 'vacancy', 'description', 'omschrijving', 'jobdesc']
                for k, v in webhook_data.items():
                    k_lower = k.lower()
                    v_str = str(v) if v else ''
                    if any(kw in k_lower for kw in vacancy_keywords) and len(v_str) > 50:
                        result['vacature'] = v_str
                        logger.info(f"📋 Found vacature in Jotform field '{k}': {len(v_str)} chars")
                        break

            # Fallback: pick the longest text field (>100 chars) as vacature
            if not result['vacature']:
                longest = max(
                    ((k, str(v)) for k, v in webhook_data.items() if v and len(str(v)) > 100),
                    key=lambda x: len(x[1]),
                    default=(None, '')
                )
                if longest[0]:
                    result['vacature'] = longest[1]
                    logger.info(f"📋 Fallback: using longest field '{longest[0]}' as vacature ({len(longest[1])} chars)")

            # Extract functie from vacature first line if not set
            if result['vacature'] and result['functie'] == 'vacature':
                result['functie'] = result['vacature'].split('\n')[0][:50]

            logger.info(f"📋 Parsed: email={result['email']}, bedrijf={result['bedrijf']}, vacature_len={len(result['vacature'])}")
            return result

        # ============================================
        # Original Typeform NESTED format
        # ============================================
        form_response = webhook_data.get('form_response', {})
        answers = form_response.get('answers', [])

        logger.info(f"📋 Parsing Typeform format: {len(answers)} answers")

        # Collect all values by type
        texts = []  # All short_text values

        for i, answer in enumerate(answers):
            if not isinstance(answer, dict):
                continue

            field = answer.get('field', {})
            if not isinstance(field, dict):
                continue

            field_type = field.get('type', '')
            field_id = field.get('id', '')

            logger.info(f"📋 Answer {i}: type={field_type}, id={field_id}")

            # Extract value based on field type
            if field_type == 'email':
                result['email'] = answer.get('email', '')
                logger.info(f"✅ Found email: {result['email']}")

            elif field_type == 'phone_number':
                result['telefoon'] = answer.get('phone_number', '')
                logger.info(f"✅ Found phone: {result['telefoon']}")

            elif field_type == 'short_text':
                text = answer.get('text', '')
                texts.append(text)
                logger.info(f"📝 Found text: {text[:50]}...")

            elif field_type == 'long_text':
                text = answer.get('text', '')
                result['vacature'] = text
                result['functie'] = text.split('\n')[0][:50] if text else 'vacature'
                logger.info(f"📝 Found long text (vacature)")

            elif field_type == 'multiple_choice':
                choice = answer.get('choice', {})
                if isinstance(choice, dict):
                    label = choice.get('label', '')
                    if not result['sector']:
                        result['sector'] = label
                    logger.info(f"📝 Found choice: {label}")

            elif field_type == 'file_upload':
                result['file_url'] = answer.get('file_url', '')
                logger.info(f"📎 Found file: {result['file_url'][:50]}...")

            elif field_type == 'contact_info':
                contact_info = answer.get('contact_info', {})
                if isinstance(contact_info, dict):
                    if contact_info.get('email'):
                        result['email'] = contact_info['email']
                    if contact_info.get('first_name'):
                        result['voornaam'] = contact_info['first_name']
                        result['contact'] = f"{contact_info.get('first_name', '')} {contact_info.get('last_name', '')}".strip()
                    if contact_info.get('phone_number'):
                        result['telefoon'] = contact_info['phone_number']
                    if contact_info.get('company'):
                        result['bedrijf'] = contact_info['company']
                    logger.info(f"✅ Found contact_info block")

        # Process collected texts (voornaam, achternaam, bedrijf order)
        if texts:
            if len(texts) >= 1 and (not result['voornaam'] or result['voornaam'] == 'daar'):
                result['voornaam'] = texts[0]
                result['contact'] = texts[0]
            if len(texts) >= 2:
                result['contact'] = f"{texts[0]} {texts[1]}".strip()
            if len(texts) >= 3 and (not result['bedrijf'] or result['bedrijf'] == 'Onbekend'):
                result['bedrijf'] = texts[2]  # 3rd text field is company

        logger.info(f"📋 Final: email={result['email']}, contact={result['contact']}, bedrijf={result['bedrijf']}")

    except Exception as e:
        logger.error(f"❌ Parse error: {e}", exc_info=True)

    return result


# ============================================================================
# LEMLIST — Automatisch lead toevoegen aan email campagne
# ============================================================================

def add_lead_to_lemlist(p: dict, analysis: dict, rapport_url: str = "") -> bool:
    """
    Voeg een lead toe aan de Lemlist email campagne na vacature-analyse.
    Vult alle custom variabelen in die de 4 email templates gebruiken.
    Returnt True bij succes, False bij mislukking of als API key ontbreekt.
    """
    if not LEMLIST_API_KEY or not LEMLIST_CAMPAIGN_ID:
        logger.warning("⚠️ Lemlist skip: LEMLIST_API_KEY of LEMLIST_CAMPAIGN_ID niet ingesteld")
        return False

    email = p.get('email', '')
    if not email or '@' not in email:
        logger.error("❌ Lemlist: geen geldig emailadres")
        return False

    try:
        improvements = analysis.get('top_3_improvements', [])
        market = analysis.get('market_analysis', {})
        salary = analysis.get('salary_benchmark', {})

        lead_data = {
            # Standaard Lemlist velden
            "firstName": p.get('voornaam', ''),
            "companyName": p.get('bedrijf', ''),
            # Custom variabelen voor de 4 email templates
            "functie": p.get('functie', ''),
            "score": str(analysis.get('overall_score', '')),
            "rapportUrl": rapport_url,
            "topImprovement1": improvements[0] if len(improvements) > 0 else '',
            "topImprovement2": improvements[1] if len(improvements) > 1 else '',
            "topImprovement3": improvements[2] if len(improvements) > 2 else '',
            "improvedText": analysis.get('improved_text', '')[:2000],  # Lemlist max
            "marketCompeting": str(market.get('competing_vacancies', '')),
            "marketCandidates": str(market.get('potential_candidates', '')),
            "salaryWarning": salary.get('warning', salary.get('difference', '')),
        }

        url = f"https://api.lemlist.com/api/campaigns/{LEMLIST_CAMPAIGN_ID}/leads/{email}"
        response = requests.post(
            url,
            auth=('', LEMLIST_API_KEY),
            json=lead_data,
            timeout=15
        )

        if response.status_code in (200, 201):
            logger.info(f"✅ Lemlist: lead {email} toegevoegd aan campagne {LEMLIST_CAMPAIGN_ID}")
            return True
        elif response.status_code == 409:
            # Lead bestaat al in deze campagne
            logger.info(f"ℹ️ Lemlist: lead {email} stond al in campagne (409)")
            return True
        else:
            logger.error(f"❌ Lemlist API fout: {response.status_code} — {response.text[:200]}")
            return False

    except Exception as e:
        logger.error(f"❌ Lemlist integratie mislukt: {e}")
        return False


def create_pipedrive_organization(name):
    """Create organization in Pipedrive"""
    if not PIPEDRIVE_API_TOKEN or not name or name == 'Onbekend':
        return None
    try:
        r = requests.post(
            f"{PIPEDRIVE_BASE}/organizations",
            params={"api_token": PIPEDRIVE_API_TOKEN},
            json={"name": name},
            timeout=30
        )
        if r.status_code == 201:
            org_id = r.json().get('data', {}).get('id')
            logger.info(f"✅ Created organization: {name} (ID: {org_id})")
            return org_id
        else:
            logger.warning(f"Org creation failed: {r.status_code} - {r.text[:200]}")
    except Exception as e:
        logger.error(f"Pipedrive org error: {e}")
    return None


def get_or_create_organization(name):
    """Get existing organization or create new one (DEDUPLICATION FIX)"""
    if not PIPEDRIVE_API_TOKEN or not name or name == 'Onbekend':
        return None

    try:
        # Search for existing organization
        r = requests.get(
            f"{PIPEDRIVE_BASE}/organizations/find",
            params={
                "api_token": PIPEDRIVE_API_TOKEN,
                "term": name,
                "limit": 1
            },
            timeout=30
        )

        if r.status_code == 200:
            data = r.json().get('data', [])
            if data and len(data) > 0:
                org_id = data[0]['id']
                logger.info(f"✅ Found existing organization: {name} (ID: {org_id})")
                return org_id

        # Not found, create new
        return create_pipedrive_organization(name)

    except Exception as e:
        logger.error(f"Error searching orgs: {e}")
        # Fallback to creation
        return create_pipedrive_organization(name)


def create_pipedrive_person(contact, email, telefoon, org_id=None):
    if not PIPEDRIVE_API_TOKEN:
        return None
    try:
        data = {
            "name": contact,
            "email": [{"value": email, "primary": True}],
            "phone": [{"value": telefoon, "primary": True}] if telefoon else []
        }
        if org_id:
            data["org_id"] = org_id

        r = requests.post(
            f"{PIPEDRIVE_BASE}/persons",
            params={"api_token": PIPEDRIVE_API_TOKEN},
            json=data,
            timeout=30
        )
        if r.status_code == 201:
            person_id = r.json().get('data', {}).get('id')
            logger.info(f"✅ Created person: {contact} (ID: {person_id})")
            return person_id
        else:
            logger.warning(f"Person creation failed: {r.status_code} - {r.text[:200]}")
    except Exception as e:
        logger.error(f"Pipedrive person error: {e}")
    return None


def create_pipedrive_deal(title, person_id, org_id=None, vacature="", file_url="", analysis="", storage_prefix=""):
    if not PIPEDRIVE_API_TOKEN:
        return None
    try:
        deal_data = {
            "title": title,
            "person_id": person_id,
            "pipeline_id": PIPELINE_ID,
            "stage_id": STAGE_ID
        }
        if org_id:
            deal_data["org_id"] = org_id
        # Store Supabase storage prefix for nurture drip emails
        if storage_prefix and FIELD_ANALYSE_STORAGE_PREFIX:
            deal_data[FIELD_ANALYSE_STORAGE_PREFIX] = storage_prefix

        r = requests.post(
            f"{PIPEDRIVE_BASE}/deals",
            params={"api_token": PIPEDRIVE_API_TOKEN},
            json=deal_data,
            timeout=30
        )
        if r.status_code == 201:
            deal_id = r.json().get('data', {}).get('id')
            logger.info(f"✅ Created deal: {title} (ID: {deal_id})")

            # Build note content
            note_parts = []
            if vacature:
                note_parts.append(f"📋 VACATURE:\n{vacature[:2000]}")
            if file_url:
                note_parts.append(f"📎 BESTAND:\n{file_url}")
            if analysis:
                note_parts.append(f"🤖 ANALYSE:\n{analysis}")

            if note_parts:
                requests.post(
                    f"{PIPEDRIVE_BASE}/notes",
                    params={"api_token": PIPEDRIVE_API_TOKEN},
                    json={
                        "deal_id": deal_id,
                        "content": "\n\n".join(note_parts)
                    },
                    timeout=30
                )
            return deal_id
        else:
            logger.warning(f"Deal creation failed: {r.status_code} - {r.text[:200]}")
    except Exception as e:
        logger.error(f"Pipedrive deal error: {e}")
    return None


@app.route("/", methods=["GET"])
def home():
    return jsonify({"status": "ok", "version": "5.3"}), 200


def retry_with_backoff(func, max_retries=3, backoff_seconds=2):
    """Retry a function with exponential backoff (ERROR RECOVERY FIX)"""
    for attempt in range(max_retries):
        try:
            return func()
        except Exception as e:
            if attempt == max_retries - 1:
                logger.error(f"❌ Failed after {max_retries} attempts: {e}")
                raise

            wait_time = backoff_seconds * (2 ** attempt)
            logger.warning(f"⏳ Attempt {attempt + 1} failed, retrying in {wait_time}s: {e}")
            time.sleep(wait_time)


def verify_jotform_signature(request_obj):
    """Verify Jotform webhook signature (SECURITY FIX)

    For testing/development: Set TEST_MODE=1 to bypass signature verification
    """
    # TESTING MODE: Allow bypassing signature check for development
    if os.getenv('TEST_MODE') == '1':
        logger.info("⚠️  TEST_MODE enabled - skipping signature verification")
        return True

    signature = request_obj.headers.get('X-Jotform-Signature')

    if not signature:
        logger.warning("❌ Missing X-Jotform-Signature header")
        return False

    payload = request_obj.get_data()
    secret = os.getenv('JOTFORM_WEBHOOK_SECRET')

    if not secret:
        logger.error("❌ JOTFORM_WEBHOOK_SECRET not configured!")
        return False

    # Jotform uses HMAC SHA256 with API key as secret
    try:
        # Compute expected signature
        expected_signature = base64.b64encode(
            hmac.new(
                secret.encode(),
                payload,
                hashlib.sha256
            ).digest()
        ).decode()

        # Compare signatures (constant-time comparison)
        is_valid = hmac.compare_digest(
            signature,
            expected_signature
        )

        if not is_valid:
            logger.error(f"❌ Invalid Jotform signature: {signature[:20]}... vs {expected_signature[:20]}...")
        else:
            logger.info("✅ Jotform webhook signature verified")

        return is_valid
    except Exception as e:
        logger.error(f"❌ Signature verification error: {e}")
        return False


def process_vacancy_analysis(p, vacancy_text):
    """
    Background task: Extract file, run Claude analysis, send email, create Pipedrive records
    This runs in a separate thread to avoid webhook timeout
    """
    try:
        logger.info(f"🔄 Background task started for {p['email']}")

        # Get vacancy text - prefer file upload over text field
        final_text = vacancy_text

        # Try to extract text from uploaded file (PDF, DOCX, DOC)
        if p['file_url']:
            logger.info(f"📎 File uploaded, attempting extraction...")
            extracted_text = extract_text_from_file(p['file_url'])
            if extracted_text and len(extracted_text) > 50:
                logger.info(f"✅ Using extracted file text ({len(extracted_text)} chars)")
                final_text = extracted_text
            else:
                logger.info(f"⚠️ File extraction failed or empty, using text field")

        # Run Claude analysis if we have vacancy text (with retry logic)
        analysis = None
        analysis_sent = False
        if final_text and len(final_text) > 50:
            try:
                # RETRY FIX: Use exponential backoff for Claude API
                analysis = retry_with_backoff(
                    lambda: analyze_vacancy_with_claude(final_text, p['bedrijf'], p['sector']),
                    max_retries=3,
                    backoff_seconds=2
                )
                if analysis:
                    # Upload analysis JSON for nurture drip emails
                    storage_prefix = ""
                    try:
                        if REPORT_BUILDER_AVAILABLE:
                            storage_prefix = upload_analysis_json(analysis, p['bedrijf'])
                            if storage_prefix:
                                logger.info(f"📦 Analysis JSON uploaded: {storage_prefix}")
                    except Exception as e:
                        logger.warning(f"⚠️ Analysis JSON upload failed: {e}")

                    try:
                        # Try hosted rapport + premium email first
                        if REPORT_BUILDER_AVAILABLE and analysis.get('categories'):
                            logger.info("📊 Building hosted rapport...")
                            rapport_html = build_hosted_rapport(analysis, p['bedrijf'], p['functie'])
                            rapport_url = upload_rapport(rapport_html, lead_name=p['bedrijf'])
                            logger.info(f"📊 Rapport URL: {rapport_url or 'upload skipped'}")

                            email_html = build_email_summary(analysis, p['bedrijf'], p['functie'], rapport_url)
                            analysis_sent = send_email(p['email'],
                                f"📊 Vacature-analyse: {p['functie']} — Score {analysis.get('overall_score', '?')}/100",
                                email_html)
                            logger.info(f"✅ Premium analysis email sent to {p['email']}")
                        else:
                            # Fallback: legacy inline email
                            analysis_sent = send_analysis_email(p['email'], p['voornaam'], p['bedrijf'], p['functie'], analysis, final_text)
                            logger.info(f"✅ Legacy analysis email sent to {p['email']}")
                    except Exception as email_error:
                        logger.error(f"❌ Report builder failed, falling back to legacy: {email_error}", exc_info=True)
                        try:
                            analysis_sent = send_analysis_email(p['email'], p['voornaam'], p['bedrijf'], p['functie'], analysis, final_text)
                        except Exception:
                            analysis_sent = False
            except Exception as e:
                logger.error(f"❌ Claude analysis failed even with retries: {e}")

        # Build analysis summary for Pipedrive notes
        analysis_summary = ""
        if analysis:
            analysis_summary = f"""SCORE: {analysis.get('overall_score', 'N/A')}/100
{analysis.get('score_section', '')}

TOP 3 VERBETERPUNTEN:
{chr(10).join(['- ' + imp for imp in analysis.get('top_3_improvements', [])])}

VERBETERDE TEKST:
{analysis.get('improved_text', '')[:1500]}"""

        # Create Pipedrive records (organization first with dedup, then person, then deal)
        org_id = get_or_create_organization(p['bedrijf'])
        person_id = create_pipedrive_person(p['contact'], p['email'], p['telefoon'], org_id)
        deal_id = create_pipedrive_deal(
            f"Vacature Analyse - {p['functie']} - {p['bedrijf']}",
            person_id,
            org_id,
            final_text,
            p['file_url'],
            analysis_summary,
            storage_prefix
        )

        # Add lead to Lemlist email sequence (only when analysis succeeded)
        lemlist_ok = False
        if analysis:
            lemlist_ok = add_lead_to_lemlist(p, analysis, rapport_url)

        logger.info(f"✅ Background task complete: org={org_id}, person={person_id}, deal={deal_id}, analysis_sent={analysis_sent}, lemlist={lemlist_ok}")

    except Exception as e:
        logger.error(f"❌ Background task failed: {e}", exc_info=True)


@app.route("/webhook/typeform", methods=["POST"])
def typeform_webhook():
    logger.info("🎯 WEBHOOK RECEIVED")

    # SECURITY FIX: Verify Jotform signature
    if not verify_jotform_signature(request):
        logger.error("❌ Jotform webhook signature invalid - rejecting")
        return jsonify({"error": "Invalid signature"}), 401

    try:
        data = request.get_json(force=True, silent=True)
        if data is None:
            # Jotform sends application/x-www-form-urlencoded
            data = request.form.to_dict(flat=True)
            logger.info(f"📥 URL-encoded keys: {list(data.keys())}")
        else:
            logger.info(f"📥 JSON keys: {list(data.keys())}")
        data = data or {}

        # Parse the data
        p = parse_typeform_data(data)

        # Validate email
        if not p['email'] or '@' not in p['email']:
            logger.error(f"❌ No email found in: {p}")
            return jsonify({"error": "No email", "parsed": p}), 400

        # Send confirmation email immediately (fast)
        confirmation_sent = send_confirmation_email(
            p['email'],
            p['voornaam'],
            p['bedrijf'],
            p['functie']
        )

        # Get vacancy text
        vacancy_text = p['vacature']

        # Queue background task to run asynchronously
        # This avoids webhook timeout (webhook can take 60+ seconds for Claude API)
        bg_thread = threading.Thread(
            target=process_vacancy_analysis,
            args=(p, vacancy_text),
            daemon=True
        )
        bg_thread.start()

        # Return immediately with 200 OK
        logger.info(f"✅ Webhook accepted, background processing started")

        return jsonify({
            "success": True,
            "confirmation_sent": confirmation_sent,
            "message": "Background processing started - analysis email will be sent shortly"
        }), 200

    except Exception as e:
        logger.error(f"❌ Error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500


def _check_admin_secret():
    """Verify admin secret header for protected endpoints."""
    secret = request.headers.get('X-Admin-Secret', '') or request.args.get('secret', '')
    if not ADMIN_SECRET or secret != ADMIN_SECRET:
        return False
    return True


@app.route("/test-email", methods=["GET"])
def test_email():
    if not _check_admin_secret():
        return jsonify({"error": "Unauthorized"}), 401
    to = request.args.get('to', 'artsrecruitin@gmail.com')
    ok = send_confirmation_email(to, "Test", "Test Bedrijf", "Test Vacature")
    return jsonify({"success": ok, "to": to}), 200 if ok else 500


@app.route("/test-async", methods=["POST"])
def test_async():
    """Test webhook WITHOUT signature verification - requires admin secret"""
    if not _check_admin_secret():
        return jsonify({"error": "Unauthorized"}), 401
    logger.info("🧪 TEST ASYNC WEBHOOK RECEIVED")

    try:
        data = request.get_json(force=True, silent=True) or {}
        p = parse_typeform_data(data)

        if not p['email'] or '@' not in p['email']:
            return jsonify({"error": "No email"}), 400

        logger.info(f"✅ Test webhook accepted for {p['email']}")

        # Send confirmation immediately
        send_confirmation_email(p['email'], p['voornaam'], p['bedrijf'], p['functie'])

        # Queue background task
        bg_thread = threading.Thread(
            target=process_vacancy_analysis,
            args=(p, p['vacature']),
            daemon=True
        )
        bg_thread.start()

        return jsonify({
            "success": True,
            "message": "Background processing started",
            "parsed": {"email": p['email'], "bedrijf": p['bedrijf'], "functie": p['functie']}
        }), 200

    except Exception as e:
        logger.error(f"❌ Error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500


@app.route("/test-pipeline", methods=["GET"])
def test_pipeline():
    """Debug: test report builder + Supabase upload (no Claude, instant)."""
    if not _check_admin_secret():
        return jsonify({"error": "Unauthorized"}), 401
    results = {
        "report_builder": REPORT_BUILDER_AVAILABLE,
        "supabase_url": bool(os.getenv("SUPABASE_URL")),
        "supabase_key": bool(os.getenv("SUPABASE_SERVICE_KEY")),
    }

    # Test with mock analysis data
    mock_analysis = {
        "overall_score": 64,
        "samenvatting": "Test rapport — pipeline verificatie",
        "categories": [
            {"name": "Vacaturetitel", "score": 75, "status": "ok"},
            {"name": "Functieomschrijving", "score": 70, "status": "ok"},
            {"name": "Salaris", "score": 35, "status": "bad"},
        ],
        "market_analysis": {"competing_vacancies": 23, "potential_candidates": 142, "market_median_salary": "€4.800", "supply_demand_ratio": "3.2x"},
        "salary_benchmark": {"offered_range": "€4.200-€5.500", "market_range": "€4.800-€6.200", "difference": "-12%", "warning": "Onder markt"},
        "improved_text": "Test verbeterde tekst",
        "action_items": ["Actie 1", "Actie 2"],
        "recommended_channels": [{"name": "Indeed", "description": "Test", "status": "AANBEVOLEN"}],
    }

    if REPORT_BUILDER_AVAILABLE:
        try:
            html = build_hosted_rapport(mock_analysis, "Test BV", "Test Functie")
            results["rapport_html_len"] = len(html)
        except Exception as e:
            results["rapport_error"] = str(e)
            return jsonify(results), 500

        try:
            url = upload_rapport(html, lead_name="test_pipeline")
            results["rapport_url"] = url
        except Exception as e:
            results["upload_error"] = str(e)

    return jsonify(results), 200


@app.route("/debug", methods=["POST"])
def debug_webhook():
    """Debug endpoint - returns what was received"""
    if not _check_admin_secret():
        return jsonify({"error": "Unauthorized"}), 401
    data = request.get_json(force=True, silent=True) or {}
    parsed = parse_typeform_data(data)
    return jsonify({
        "received_keys": list(data.keys()),
        "parsed": parsed,
    }), 200


# =============================================================================
# TRUST-FIRST EMAIL NURTURE SYSTEM V5.1
# =============================================================================

def get_nurture_email_html(email_num, voornaam, functie_titel, analysis_data=None):
    """Generate HTML content for nurture emails.

    Emails 1-3 are VALUE DRIP emails that deliver analysis content:
    - Email 1: Verbeterde vacaturetekst (the golden content)
    - Email 2: Marktanalyse + salaris benchmark
    - Email 3: Kanaalstrategie + actieplan
    Emails 4-8 are trust-building tips and check-ins (unchanged).
    """
    ad = analysis_data or {}

    # Email 1: VERBETERDE VACATURETEKST
    improved_text = ad.get('improved_text', '')
    improved_html = improved_text.replace('\n', '<br>') if improved_text else ''
    top_improvements = ad.get('top_3_improvements', [])
    improvements_list = ''.join(f'<li style="margin-bottom:8px;">{imp}</li>' for imp in top_improvements) if top_improvements else ''

    # Email 2: MARKTANALYSE + SALARIS
    ma = ad.get('market_analysis', {})
    sb = ad.get('salary_benchmark', {})

    # Email 3: KANAALSTRATEGIE + ACTIEPLAN
    channels = ad.get('recommended_channels', [])
    channels_html = ''.join(
        f'<tr><td style="padding:10px;border-bottom:1px solid #eee;font-weight:bold;color:#1f2937;">{ch.get("name","")}</td>'
        f'<td style="padding:10px;border-bottom:1px solid #eee;color:#6b7280;">{ch.get("description","")}</td>'
        f'<td style="padding:10px;border-bottom:1px solid #eee;"><span style="background:#dcfce7;color:#166534;padding:3px 8px;border-radius:3px;font-size:11px;font-weight:bold;">{ch.get("status","")}</span></td></tr>'
        for ch in channels
    ) if channels else ''
    action_items = ad.get('action_items', [])
    actions_html = ''.join(
        f'<tr><td style="padding:10px 0;border-bottom:1px solid #f3f4f6;">'
        f'<table cellpadding="0" cellspacing="0"><tr>'
        f'<td style="width:28px;height:28px;background:#f59e0b;color:white;text-align:center;font-weight:bold;font-size:13px;line-height:28px;border-radius:50%;">{i+1}</td>'
        f'<td style="padding-left:12px;color:#1f2937;font-size:14px;">{item}</td>'
        f'</tr></table></td></tr>'
        for i, item in enumerate(action_items)
    ) if action_items else ''

    templates = {
        1: f"""<div style="font-family: Arial, sans-serif; font-size: 14px; line-height: 1.6; color: #333333;">
<p>Hoi {voornaam},</p>
<p>Gisteren ontving je de score-analyse van je vacature voor <strong>{functie_titel}</strong>. Vandaag het belangrijkste onderdeel: <strong>de verbeterde vacaturetekst</strong>.</p>
<p>Onze AI heeft je tekst herschreven met 3 concrete verbeteringen:</p>
{f'<ul style="margin:15px 0;padding-left:20px;background:#fffbeb;border-left:4px solid #f59e0b;padding:15px 15px 15px 35px;">{improvements_list}</ul>' if improvements_list else ''}
<div style="background:#ecfdf5;border:2px solid #10b981;padding:20px;margin:20px 0;border-radius:6px;">
<div style="background:#10b981;color:white;display:inline-block;padding:4px 12px;border-radius:3px;font-size:11px;font-weight:bold;margin-bottom:12px;">VERBETERDE TEKST — DIRECT TE GEBRUIKEN</div>
<div style="background:white;border:1px solid #a7f3d0;padding:16px;font-size:13px;color:#374151;line-height:22px;margin-top:8px;">
{improved_html if improved_html else '<em>Verbeterde tekst is beschikbaar in je volledige rapport.</em>'}
</div>
</div>
<p><strong>Tip:</strong> Kopieer de tekst hierboven en plaats hem direct op je vacatureplatform. De meeste recruiters zien binnen 48 uur verschil in response.</p>
<p>Morgen stuur ik je de marktanalyse — hoe jouw vacature zich verhoudt tot de concurrentie.</p>
<p>Groeten,<br><strong>Wouter</strong><br><span style="color: #666666;">kandidatentekort.nl</span></p>
</div>""",

        2: f"""<div style="font-family: Arial, sans-serif; font-size: 14px; line-height: 1.6; color: #333333;">
<p>Hoi {voornaam},</p>
<p>Gisteren de verbeterde tekst, vandaag de data: <strong>hoe staat jouw {functie_titel} vacature ervoor op de arbeidsmarkt?</strong></p>
<div style="background:#f8fafc;border:1px solid #e2e8f0;padding:20px;margin:20px 0;border-radius:6px;">
<div style="font-size:16px;font-weight:bold;color:#1e3a8a;margin-bottom:15px;">📊 Marktanalyse</div>
<table width="100%" cellpadding="0" cellspacing="0" style="margin-bottom:15px;">
<tr>
<td width="50%" style="padding:12px;background:#f1f5f9;text-align:center;border-radius:4px;">
<div style="font-size:28px;font-weight:bold;color:#1f2937;">{ma.get('competing_vacancies', '—')}</div>
<div style="font-size:11px;color:#6b7280;margin-top:4px;">concurrerende vacatures</div>
</td>
<td width="8px"></td>
<td width="50%" style="padding:12px;background:#f1f5f9;text-align:center;border-radius:4px;">
<div style="font-size:28px;font-weight:bold;color:#1f2937;">{ma.get('potential_candidates', '—')}</div>
<div style="font-size:11px;color:#6b7280;margin-top:4px;">potentiële kandidaten</div>
</td>
</tr>
</table>
<table width="100%" cellpadding="0" cellspacing="0">
<tr>
<td width="50%" style="padding:12px;background:#f1f5f9;text-align:center;border-radius:4px;">
<div style="font-size:28px;font-weight:bold;color:#1f2937;">{ma.get('market_median_salary', '—')}</div>
<div style="font-size:11px;color:#6b7280;margin-top:4px;">markt mediaan salaris</div>
</td>
<td width="8px"></td>
<td width="50%" style="padding:12px;background:#f1f5f9;text-align:center;border-radius:4px;">
<div style="font-size:28px;font-weight:bold;color:#1f2937;">{ma.get('supply_demand_ratio', '—')}</div>
<div style="font-size:11px;color:#6b7280;margin-top:4px;">vraag/aanbod ratio</div>
</td>
</tr>
</table>
</div>
<div style="background:#fef2f2;border-left:4px solid #ef4444;padding:15px 20px;margin:20px 0;border-radius:0 6px 6px 0;">
<div style="font-weight:bold;color:#991b1b;margin-bottom:6px;">💰 Salaris Benchmark</div>
<table width="100%" cellpadding="0" cellspacing="0" style="font-size:13px;">
<tr><td style="color:#6b7280;padding:4px 0;">Aangeboden:</td><td style="font-weight:bold;color:#1f2937;">{sb.get('offered_range', 'niet vermeld')}</td></tr>
<tr><td style="color:#6b7280;padding:4px 0;">Markt range:</td><td style="font-weight:bold;color:#1f2937;">{sb.get('market_range', '—')}</td></tr>
<tr><td style="color:#6b7280;padding:4px 0;">Verschil:</td><td style="font-weight:bold;color:#ef4444;">{sb.get('difference', '—')}</td></tr>
</table>
{f'<p style="margin:10px 0 0;font-size:12px;color:#991b1b;">⚠️ {sb.get("warning", "")}</p>' if sb.get('warning') else ''}
</div>
<p><strong>Wat betekent dit?</strong> Als je vraag/aanbod ratio hoger is dan 2x, moet je vacature écht opvallen om de juiste kandidaten te trekken. De verbeterde tekst van gisteren helpt daarbij.</p>
<p>Overmorgen stuur ik het laatste deel: je <strong>kanaalstrategie en actieplan</strong>.</p>
<p>Groeten,<br><strong>Wouter</strong></p>
</div>""",

        3: f"""<div style="font-family: Arial, sans-serif; font-size: 14px; line-height: 1.6; color: #333333;">
<p>Hoi {voornaam},</p>
<p>Laatste deel van je analyse voor <strong>{functie_titel}</strong>: waar moet je je vacature plaatsen en wat zijn je volgende stappen?</p>
{f'''<div style="background:#f8fafc;border:1px solid #e2e8f0;padding:20px;margin:20px 0;border-radius:6px;">
<div style="font-size:16px;font-weight:bold;color:#1e3a8a;margin-bottom:15px;">📡 Aanbevolen Kanalen</div>
<table width="100%" cellpadding="0" cellspacing="0" style="font-size:13px;">
<tr style="background:#f1f5f9;"><th style="padding:10px;text-align:left;color:#6b7280;font-size:11px;">KANAAL</th><th style="padding:10px;text-align:left;color:#6b7280;font-size:11px;">WAAROM</th><th style="padding:10px;text-align:left;color:#6b7280;font-size:11px;">STATUS</th></tr>
{channels_html}
</table>
</div>''' if channels_html else ''}
{f'''<div style="background:#fffbeb;border:2px solid #f59e0b;padding:20px;margin:20px 0;border-radius:6px;">
<div style="font-size:16px;font-weight:bold;color:#92400e;margin-bottom:15px;">🎯 Actieplan — 5 stappen</div>
<table width="100%" cellpadding="0" cellspacing="0">{actions_html}</table>
</div>''' if actions_html else ''}
<div style="background:#f0f4ff;border-left:4px solid #3b82f6;padding:15px 20px;margin:20px 0;">
<p style="margin:0 0 8px;font-weight:bold;color:#1e40af;">Samenvatting van je 3 emails:</p>
<ol style="margin:0;padding-left:20px;color:#374151;">
<li><strong>Dag 1:</strong> Verbeterde vacaturetekst (direct toepasbaar)</li>
<li><strong>Dag 3:</strong> Marktanalyse + salaris benchmark (vandaag)</li>
<li><strong>Dag 5:</strong> Kanaalstrategie + actieplan (deze email)</li>
</ol>
</div>
<p>Dit was de complete analyse. Heb je vragen over een van de onderdelen? Reply gewoon — ik lees alles persoonlijk.</p>
<p>De komende weken stuur ik je nog een paar recruitment tips die specifiek relevant zijn voor jouw sector.</p>
<p>Succes met de werving,<br><strong>Wouter</strong></p>
</div>""",

        4: f"""<div style="font-family: Arial, sans-serif; font-size: 14px; line-height: 1.6; color: #333333;">
<p>Hoi {voornaam},</p>
<p>Deze week deel ik een tip die veel recruiters over het hoofd zien:</p>
<div style="background-color: #FEF3C7; border-left: 4px solid #F59E0B; padding: 15px 20px; margin: 20px 0;">
<p style="margin: 0 0 10px 0; font-size: 16px;"><strong>De functietitel bepaalt 70% van je zichtbaarheid</strong></p>
<p style="margin: 0 0 15px 0;">Kandidaten zoeken op specifieke termen. Een creatieve titel als "Teamspeler Extraordinaire" klinkt leuk, maar niemand zoekt daarop.</p>
<p style="margin: 0 0 10px 0;"><strong>Wat werkt:</strong></p>
<ul style="margin: 0 0 15px 0; padding-left: 20px;">
<li>Gebruik de exacte term die kandidaten googlen</li>
<li>Voeg niveau toe (Junior/Medior/Senior)</li>
<li>Houd het onder 60 karakters</li>
</ul>
<p style="margin: 0;"><strong>Voorbeeld:</strong><br>"Commerciele Binnendienst Medewerker" krijgt 3x meer views dan "Sales Ninja"</p>
</div>
<p>Heb je al gedacht aan het A/B testen van je functietitels?</p>
<p>Succes deze week,<br><strong>Wouter</strong></p>
</div>""",

        5: f"""<div style="font-family: Arial, sans-serif; font-size: 14px; line-height: 1.6; color: #333333;">
<p>Hoi {voornaam},</p>
<p>"Salaris: marktconform" - de meest waardeloze zin in recruitment.</p>
<p>Hier is wat de data zegt:</p>
<div style="background-color: #D1FAE5; border-left: 4px solid #10B981; padding: 15px 20px; margin: 20px 0;">
<p style="margin: 0 0 10px 0;"><strong>Vacatures met salarisindicatie krijgen:</strong></p>
<ul style="margin: 0 0 10px 0; padding-left: 20px; list-style: none;">
<li>✅ +35% meer sollicitaties</li>
<li>✅ +27% hogere kwaliteit kandidaten</li>
<li>✅ +40% snellere time-to-hire</li>
</ul>
<p style="margin: 0; font-size: 12px; color: #666; font-style: italic;">Bron: Indeed Hiring Lab 2024</p>
</div>
<p><strong>Maar wat als je het niet mag vermelden?</strong></p>
<p>Alternatieven die ook werken:</p>
<ul style="margin: 15px 0; padding-left: 20px;">
<li>"Salarisindicatie: vanaf EUR 3.500 bruto/maand"</li>
<li>"Indicatie: schaal 8-10 CAO [naam]"</li>
<li>"Budget: EUR 45.000 - 55.000 op jaarbasis"</li>
</ul>
<p>Zelfs een range is beter dan niets.</p>
<p>Groeten,<br><strong>Wouter</strong></p>
</div>""",

        6: f"""<div style="font-family: Arial, sans-serif; font-size: 14px; line-height: 1.6; color: #333333;">
<p>Hoi {voornaam},</p>
<p>Wist je dat kandidaten gemiddeld <strong>6 seconden</strong> besteden aan de eerste scan van een vacature?</p>
<p>In die 6 seconden beslissen ze of ze doorlezen of wegklikken.</p>
<div style="background-color: #EDE9FE; border-left: 4px solid #7C3AED; padding: 15px 20px; margin: 20px 0;">
<p style="margin: 0 0 10px 0;"><strong>Wat ze scannen:</strong></p>
<ol style="margin: 0 0 15px 0; padding-left: 20px;">
<li>Functietitel</li>
<li>Salaris (als het er staat)</li>
<li>De eerste 2-3 zinnen</li>
<li>Locatie</li>
<li>Logo/bedrijfsnaam</li>
</ol>
<p style="margin: 0 0 10px 0;"><strong>Het probleem:</strong> 90% begint met "Wij zoeken een enthousiaste..."</p>
<p style="margin: 0;"><strong>De oplossing:</strong> Start met een vraag of bold statement.</p>
</div>
<p>Pak eens een van je huidige vacatures erbij. Hoe is de opening?</p>
<p>Tot volgende week,<br><strong>Wouter</strong></p>
<p style="color: #999999; font-size: 12px; margin-top: 30px; padding-top: 15px; border-top: 1px solid #eeeeee;">
Dit was de laatste tip in deze serie. Vond je ze nuttig? Laat het me weten.</p>
</div>""",

        7: f"""<div style="font-family: Arial, sans-serif; font-size: 14px; line-height: 1.6; color: #333333;">
<p>Hoi {voornaam},</p>
<p>Het is nu drie weken geleden dat je de vacature-analyse ontving.</p>
<p>Ik ben benieuwd hoe het gaat met je werving. Heb je de kandidaat al gevonden? Of loop je nog ergens tegenaan?</p>
<div style="background-color: #f8f9fa; border-left: 4px solid #EF7D00; padding: 15px 20px; margin: 20px 0;">
<p style="margin: 0 0 10px 0; font-size: 16px;"><strong>Zullen we even bellen?</strong></p>
<p style="margin: 0 0 15px 0;">Geen verkooppraatje, gewoon een kort gesprek (15 min) om te kijken of ik je ergens mee kan helpen.</p>
<p style="margin: 0 0 10px 0;"><strong>We kunnen het hebben over:</strong></p>
<ul style="margin: 0 0 15px 0; padding-left: 20px;">
<li>De resultaten van je huidige vacature</li>
<li>Andere openstaande posities</li>
<li>Recruitment uitdagingen waar je tegenaan loopt</li>
</ul>
<p style="margin: 0;">
<a href="https://calendly.com/wouter-arts-/vacature-analyse-advies" style="display: inline-block; background-color: #EF7D00; color: #ffffff; padding: 12px 25px; text-decoration: none; border-radius: 5px; font-weight: bold;">Plan een moment dat jou uitkomt</a>
</p>
</div>
<p>Geen zin of geen tijd? Reply dan gewoon even met een update.</p>
<p>Groeten,<br><strong>Wouter</strong></p>
</div>""",

        8: f"""<div style="font-family: Arial, sans-serif; font-size: 14px; line-height: 1.6; color: #333333;">
<p>Hoi {voornaam},</p>
<p>Een maand geleden verstuurde ik de analyse voor je <strong>{functie_titel}</strong> vacature.</p>
<p>Dit is mijn laatste mail in deze serie - daarna laat ik je met rust.</p>
<p>Maar voor ik ga, ben ik nieuwsgierig:</p>
<div style="background-color: #f8f9fa; border: 1px solid #e0e0e0; padding: 15px 20px; margin: 20px 0; border-radius: 5px;">
<p style="margin: 0 0 10px 0;"><strong>Hoe is het gegaan?</strong></p>
<table style="width: 100%; border-collapse: collapse;">
<tr><td style="padding: 5px 0;"><strong>A.</strong> Kandidaat gevonden - top!</td></tr>
<tr><td style="padding: 5px 0;"><strong>B.</strong> Nog bezig - maar gaat goed</td></tr>
<tr><td style="padding: 5px 0;"><strong>C.</strong> Vacature on hold gezet</td></tr>
<tr><td style="padding: 5px 0;"><strong>D.</strong> Hulp nodig - laten we bellen</td></tr>
</table>
<p style="margin: 10px 0 0 0; font-size: 13px; color: #666;">Reply met A, B, C of D (of vertel gewoon je verhaal)</p>
</div>
<p>Hoe dan ook - succes met je recruitment!</p>
<p>Groeten,<br><strong>Wouter</strong><br><span style="color: #666666;">kandidatentekort.nl</span></p>
<p style="color: #999999; font-size: 12px; margin-top: 30px; padding-top: 15px; border-top: 1px solid #eeeeee;">
Contact: warts@recruitin.nl | Nieuwe vacature? <a href="https://kandidatentekort.nl" style="color: #EF7D00;">kandidatentekort.nl</a></p>
</div>"""
    }

    return templates.get(email_num, "")


def get_nurture_email_subject(email_num, functie_titel="vacature"):
    """Get subject line for nurture email"""
    subjects = {
        1: f"Je verbeterde vacaturetekst voor {functie_titel}",
        2: f"Marktanalyse: {functie_titel} — hoe staat jouw vacature ervoor?",
        3: f"Kanaalstrategie + actieplan voor {functie_titel}",
        4: "Recruitment tip: De kracht van de juiste functietitel",
        5: "Recruitment tip: Het salarisvraagstuk",
        6: "Recruitment tip: De eerste 6 seconden",
        7: "Zullen we eens bellen?",
        8: "Laatste check - hoe staat het ervoor?"
    }
    return subjects.get(email_num, "Follow-up van kandidatentekort.nl")


def send_nurture_email(to_email, email_num, voornaam, functie_titel, analysis_data=None):
    """Send a nurture sequence email. Emails 1-3 include analysis data if available."""
    if not GMAIL_APP_PASSWORD:
        logger.warning("No Gmail password configured")
        return False

    try:
        subject = get_nurture_email_subject(email_num, functie_titel)
        html_content = get_nurture_email_html(email_num, voornaam, functie_titel, analysis_data)

        if not html_content:
            logger.error(f"No template for email {email_num}")
            return False

        msg = MIMEMultipart('alternative')
        msg['Subject'] = subject
        msg['From'] = f"Wouter van kandidatentekort.nl <{GMAIL_USER}>"
        msg['To'] = to_email
        msg['Reply-To'] = "warts@recruitin.nl"

        msg.attach(MIMEText(html_content, 'html'))

        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
            server.login(GMAIL_USER, GMAIL_APP_PASSWORD)
            server.send_message(msg)

        logger.info(f"✅ Sent nurture email {email_num} to {to_email}")
        return True

    except Exception as e:
        logger.error(f"❌ Failed to send nurture email: {e}")
        return False


def update_deal_nurture_status(deal_id, email_num):
    """Update deal fields after sending nurture email"""
    if not PIPEDRIVE_API_TOKEN:
        return False

    try:
        # Map email number to option ID (Pipedrive enum options)
        email_options = {
            1: "Email 1", 2: "Email 2", 3: "Email 3", 4: "Email 4",
            5: "Email 5", 6: "Email 6", 7: "Email 7", 8: "Email 8"
        }

        update_data = {
            FIELD_LAATSTE_EMAIL: email_options.get(email_num, "Email 1")
        }

        # If email 8, mark sequence as complete
        if email_num == 8:
            update_data[FIELD_EMAIL_SEQUENCE_STATUS] = "Completed"

        response = requests.put(
            f"{PIPEDRIVE_BASE}/deals/{deal_id}",
            params={"api_token": PIPEDRIVE_API_TOKEN},
            json=update_data,
            timeout=30
        )

        if response.status_code == 200:
            logger.info(f"✅ Updated deal {deal_id} nurture status to Email {email_num}")
            return True
        else:
            logger.warning(f"Failed to update deal: {response.status_code}")
            return False

    except Exception as e:
        logger.error(f"Error updating deal: {e}")
        return False


def get_deals_for_nurture():
    """Get all deals that need nurture emails today"""
    if not PIPEDRIVE_API_TOKEN:
        return []

    try:
        # Get all deals from pipeline
        response = requests.get(
            f"{PIPEDRIVE_BASE}/deals",
            params={
                "api_token": PIPEDRIVE_API_TOKEN,
                "pipeline_id": PIPELINE_ID,
                "status": "open",
                "limit": 500
            },
            timeout=30
        )

        if response.status_code != 200:
            logger.error(f"Failed to get deals: {response.status_code}")
            return []

        deals = response.json().get('data', []) or []
        deals_to_email = []
        today = datetime.now().date()

        for deal in deals:
            # Only process deals in Gekwalificeerd stage (21)
            stage_id = deal.get('stage_id')
            if stage_id != NURTURE_ACTIVE_STAGE:
                continue

            # Get custom field values
            rapport_date_str = deal.get(FIELD_RAPPORT_VERZONDEN)
            sequence_status = deal.get(FIELD_EMAIL_SEQUENCE_STATUS, '')
            laatste_email = deal.get(FIELD_LAATSTE_EMAIL, '')

            # Skip if no rapport date or sequence not active
            if not rapport_date_str:
                continue

            # Skip if sequence is completed, paused, or responded
            if sequence_status in ['Completed', 'Gepauzeerd', 'Voltooid', 'Responded', 'Unsubscribed']:
                continue

            # Parse rapport date
            try:
                rapport_date = datetime.strptime(rapport_date_str, '%Y-%m-%d').date()
            except (ValueError, TypeError):
                continue

            days_since_rapport = (today - rapport_date).days

            # Determine which email to send
            current_email = 0
            if laatste_email:
                try:
                    current_email = int(laatste_email.replace('Email ', ''))
                except (ValueError, TypeError):
                    pass

            next_email = current_email + 1

            # Check if it's time to send the next email
            if next_email <= 8:
                scheduled_day = EMAIL_SCHEDULE.get(next_email, {}).get('day', 999)
                if days_since_rapport >= scheduled_day:
                    # Get person info for email
                    person_id = deal.get('person_id', {})
                    if isinstance(person_id, dict):
                        person_id = person_id.get('value')

                    deals_to_email.append({
                        'deal_id': deal.get('id'),
                        'deal_title': deal.get('title', ''),
                        'person_id': person_id,
                        'next_email': next_email,
                        'days_since': days_since_rapport,
                        'storage_prefix': deal.get(FIELD_ANALYSE_STORAGE_PREFIX, '') if FIELD_ANALYSE_STORAGE_PREFIX else '',
                    })

        logger.info(f"📧 Found {len(deals_to_email)} deals in stage {NURTURE_ACTIVE_STAGE} (Gekwalificeerd) ready for nurture emails")
        return deals_to_email

    except Exception as e:
        logger.error(f"Error getting deals for nurture: {e}")
        return []


def get_person_email(person_id):
    """Get person's email and name from Pipedrive"""
    if not PIPEDRIVE_API_TOKEN or not person_id:
        return None, None

    try:
        response = requests.get(
            f"{PIPEDRIVE_BASE}/persons/{person_id}",
            params={"api_token": PIPEDRIVE_API_TOKEN},
            timeout=30
        )

        if response.status_code == 200:
            person = response.json().get('data', {})
            emails = person.get('email', [])
            email = emails[0].get('value') if emails else None
            name = person.get('first_name', 'daar')
            return email, name
    except Exception as e:
        logger.error(f"Error getting person: {e}")

    return None, None


def should_send_email(deal, email_num):
    """Check if this email should be sent (not already sent, not too early) (STATE TRACKING FIX)"""
    try:
        # Get custom field values from Pipedrive
        r = requests.get(
            f"{PIPEDRIVE_BASE}/deals/{deal['id']}",
            params={"api_token": PIPEDRIVE_API_TOKEN},
            timeout=30
        )
        deal_data = r.json().get('data', {})

        # Check if this specific email was already sent
        sent_emails = deal_data.get(FIELD_EMAIL_SEQUENCE_STATUS, '').split(',')
        if str(email_num) in sent_emails:
            logger.info(f"⏭️ Deal {deal['id']}: Email {email_num} already sent")
            return False

        # Check timing
        rapport_date = deal_data.get(FIELD_RAPPORT_VERZONDEN)
        if rapport_date:
            try:
                from dateutil import parser
                rapport_dt = parser.parse(rapport_date)
            except (ImportError, ValueError):
                rapport_dt = datetime.fromisoformat(rapport_date.split('T')[0])

            days_since = (datetime.now() - rapport_dt).days
            required_days = EMAIL_SCHEDULE[email_num]['day']

            if days_since < required_days:
                logger.info(f"⏳ Deal {deal['id']}: Email {email_num} not due yet ({days_since}/{required_days} days)")
                return False

        return True

    except Exception as e:
        logger.error(f"❌ Error checking email status: {e}")
        return False


def mark_email_sent(deal_id, email_num):
    """Record that an email was sent (STATE TRACKING FIX)"""
    try:
        # Get current status
        r = requests.get(
            f"{PIPEDRIVE_BASE}/deals/{deal_id}",
            params={"api_token": PIPEDRIVE_API_TOKEN},
            timeout=30
        )
        current_status = r.json().get('data', {}).get(FIELD_EMAIL_SEQUENCE_STATUS, '')

        # Add this email number
        sent = set(current_status.split(',')) if current_status else set()
        sent.add(str(email_num))
        new_status = ','.join(sorted(sent))

        # Update deal
        requests.patch(
            f"{PIPEDRIVE_BASE}/deals/{deal_id}",
            params={"api_token": PIPEDRIVE_API_TOKEN},
            json={
                FIELD_EMAIL_SEQUENCE_STATUS: new_status,
                FIELD_LAATSTE_EMAIL: email_num
            },
            timeout=30
        )
        logger.info(f"✅ Deal {deal_id}: Marked email {email_num} as sent")
    except Exception as e:
        logger.error(f"❌ Error marking email sent: {e}")


def process_nurture_emails():
    """Process all pending nurture emails (UPDATED with state tracking)"""
    logger.info("🔄 Starting nurture email processing...")

    deals = get_deals_for_nurture()
    sent_count = 0

    for deal in deals:
        try:
            email, voornaam = get_person_email(deal['person_id'])

            if not email:
                logger.warning(f"No email for deal {deal['deal_id']}")
                continue

            # STATE TRACKING FIX: Check if this email should be sent
            if not should_send_email(deal, deal['next_email']):
                continue

            # Extract functie from deal title
            functie_titel = deal['deal_title'].replace('Vacature Analyse - ', '').split(' - ')[0]

            # For emails 1-3 (value drip), fetch analysis data from Supabase
            analysis_data = None
            if deal['next_email'] <= 3 and REPORT_BUILDER_AVAILABLE:
                storage_prefix = deal.get('storage_prefix', '')
                if storage_prefix:
                    try:
                        analysis_data = fetch_analysis_json(storage_prefix)
                        if analysis_data:
                            logger.info(f"📦 Loaded analysis data for drip email {deal['next_email']}")
                    except Exception as e:
                        logger.warning(f"⚠️ Could not fetch analysis data: {e}")

            # Send the email
            success = send_nurture_email(
                email,
                deal['next_email'],
                voornaam or 'daar',
                functie_titel,
                analysis_data
            )

            if success:
                # STATE TRACKING FIX: Mark email as sent
                mark_email_sent(deal['deal_id'], deal['next_email'])
                sent_count += 1

            # Small delay between emails
            time.sleep(2)

        except Exception as e:
            logger.error(f"Error processing deal {deal.get('deal_id')}: {e}")

    logger.info(f"✅ Nurture processing complete: {sent_count}/{len(deals)} emails sent")
    return sent_count


def start_nurture_deal(deal_id):
    """Start nurture sequence for a specific deal"""
    if not PIPEDRIVE_API_TOKEN:
        return False

    try:
        today = datetime.now().strftime('%Y-%m-%d')

        response = requests.put(
            f"{PIPEDRIVE_BASE}/deals/{deal_id}",
            params={"api_token": PIPEDRIVE_API_TOKEN},
            json={
                FIELD_RAPPORT_VERZONDEN: today,
                FIELD_EMAIL_SEQUENCE_STATUS: "Actief"
            },
            timeout=30
        )

        if response.status_code == 200:
            logger.info(f"✅ Started nurture sequence for deal {deal_id}")
            return True
        else:
            logger.warning(f"Failed to start nurture: {response.status_code}")
            return False

    except Exception as e:
        logger.error(f"Error starting nurture: {e}")
        return False


# Background scheduler for nurture emails
def nurture_scheduler():
    """Background thread that checks for nurture emails every hour (FIXED: prevents duplicate sends)"""
    last_run_hour = -1  # Track which hour we already ran

    while True:
        try:
            # Run at specific times (9 AM, 2 PM Dutch time)
            now = datetime.now()
            current_hour = now.hour

            # Only run once per hour, at 9 AM and 2 PM
            if current_hour in [9, 14] and current_hour != last_run_hour:
                logger.info(f"⏰ Scheduled nurture check running at {now.strftime('%H:%M:%S')}")
                process_nurture_emails()
                last_run_hour = current_hour

            # Reset at midnight
            if current_hour == 0:
                last_run_hour = -1

        except Exception as e:
            logger.error(f"Scheduler error: {e}", exc_info=True)

        # Check every 60 seconds (more responsive, better accuracy)
        time.sleep(60)


@app.route("/rapport", methods=["GET"])
def serve_rapport():
    """Proxy: fetch rapport HTML from Supabase Storage and serve as text/html."""
    path = request.args.get("path", "")
    if not path:
        return "Missing 'path' parameter", 400

    # SECURITY: Validate path to prevent traversal attacks
    if not re.match(r'^\d{8}/[a-zA-Z0-9_\-]+/rapport\.html$', path):
        return "Invalid path format", 400

    supabase_url = os.getenv("SUPABASE_URL", "")
    supabase_key = os.getenv("SUPABASE_SERVICE_KEY", "")
    if not supabase_url or not supabase_key:
        return "Storage not configured", 500

    try:
        resp = requests.get(
            f"{supabase_url}/storage/v1/object/public/kt-assets/{path}",
            timeout=10,
        )
        if resp.status_code == 200:
            return resp.content, 200, {"Content-Type": "text/html; charset=utf-8", "Cache-Control": "public, max-age=3600"}
        else:
            logger.error(f"Rapport fetch failed: {resp.status_code}")
            return "Rapport niet gevonden", 404
    except Exception as e:
        logger.error(f"Rapport proxy error: {e}")
        return "Fout bij ophalen rapport", 500


@app.route("/health", methods=["GET"])
def health_check():
    return jsonify({
        "status": "healthy",
        "version": "5.3",
        "report_builder": REPORT_BUILDER_AVAILABLE,
        "features": ["typeform", "analysis", "nurture"],
        "email": bool(GMAIL_APP_PASSWORD),
        "pipedrive": bool(PIPEDRIVE_API_TOKEN),
        "claude": bool(ANTHROPIC_API_KEY)
    }), 200


@app.route("/nurture/process", methods=["POST"])
def trigger_nurture_processing():
    """Manually trigger nurture email processing"""
    if not _check_admin_secret():
        return jsonify({"error": "Unauthorized"}), 401
    count = process_nurture_emails()
    return jsonify({"success": True, "emails_sent": count}), 200


@app.route("/nurture/start/<int:deal_id>", methods=["POST"])
def start_nurture_for_deal(deal_id):
    """Start nurture sequence for a specific deal"""
    if not _check_admin_secret():
        return jsonify({"error": "Unauthorized"}), 401
    success = start_nurture_deal(deal_id)
    return jsonify({"success": success, "deal_id": deal_id}), 200 if success else 500


@app.route("/nurture/status", methods=["GET"])
def nurture_status():
    """Get status of nurture sequences"""
    if not _check_admin_secret():
        return jsonify({"error": "Unauthorized"}), 401
    deals = get_deals_for_nurture()
    return jsonify({"pending_emails": len(deals), "deals": deals[:20]}), 200


@app.route("/nurture/test/<int:email_num>", methods=["GET"])
def test_nurture_email(email_num):
    """Send a test nurture email"""
    if not _check_admin_secret():
        return jsonify({"error": "Unauthorized"}), 401
    to = request.args.get('to', 'warts@recruitin.nl')
    voornaam = request.args.get('name', 'Test')
    functie = request.args.get('functie', 'Senior Developer')
    success = send_nurture_email(to, email_num, voornaam, functie)
    return jsonify({"success": success, "email_num": email_num, "to": to}), 200 if success else 500


if __name__ == "__main__":
    # Start background scheduler for nurture emails
    scheduler_thread = threading.Thread(target=nurture_scheduler, daemon=True)
    scheduler_thread.start()
    logger.info("🚀 Nurture scheduler started")

    app.run(host="0.0.0.0", port=int(os.getenv("PORT", 8080)))
